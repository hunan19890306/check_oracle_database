# -*- coding: utf-8 -*-
from docx.enum.style import WD_STYLE_TYPE
from docx import *

document = Document()
styles = document.styles
para = document.add_paragraph()

#生成所有字符样式
for s in styles:
    if s.type == WD_STYLE_TYPE.CHARACTER:
        run = para.add_run("Character style is:  "+s.name+"\n")
        run.style = s

document.save('character_style.docx')