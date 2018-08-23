# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import cx_Oracle
import paramiko
import sys
#数据库连接
def connect_db(userinfo,host,service_name,sql):
    conn_str = userinfo + host + service_name
    print conn_str
    conn = cx_Oracle.connect(conn_str)    
    c = conn.cursor()                                       
    c.execute(sql)
    if sql.startswith('select'):
        fir=c.fetchall()
        print fir
        conn.commit()             
        c.close()                                 
        conn.close()
        return fir
    else:
        conn.commit()             
        c.close()                                 
        conn.close()
#word输出结构
def new_doc(userinfo,host,db_name,tablespace_sql,doc_path):#新建一个word文档,写入汇总表的数据
    document = Document()
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'测试')
    document.save(doc_path)
    get_tablespace(userinfo,host,db_name,tablespace_sql)
#输出表空间使用率
def get_tablespace(userinfo,host,service_name,sql):
    document = Document(doc_path)
    names = locals()
    for q in service_name:
        if q.strip():
            tablespace = connect_db(userinfo,host,q,sql)
            print tablespace
            print q
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s表空间使用率：" % q)
            r_total.font.bold = True
            names["tablespace_output%s" % q]= document.add_table(1,6,style="Light List Accent 5")
            tablespace_output_cells = names["tablespace_output%s" % q].rows[0].cells
            tablespace_output_cells[0].text = u'表空间名'
            tablespace_output_cells[1].text = u'总空间'
            tablespace_output_cells[2].text = u'已使用'
            tablespace_output_cells[3].text = u'使用率'
            tablespace_output_cells[4].text = u'剩余空间'
            tablespace_output_cells[5].text = u'最大块'
            for i in range(0,len(tablespace)):
                cells = names["tablespace_output%s" % q].add_row().cells
                cells[0].text = str(tablespace[i][0])
                cells[1].text = str(tablespace[i][1])
                cells[2].text = str(tablespace[i][2])
                cells[3].text = str(tablespace[i][3])
                cells[4].text = str(tablespace[i][4])
                cells[5].text = str(tablespace[i][5])
            print q
        else:
            pass
    document.save(doc_path)
#主函数
if __name__ == "__main__":
#连接信息
    userinfo = 'usr_dump/usr_dump'
    host = '@172.16.9.75:1521/'
#ssh连接到节点一取的信息
    ssh_db1 = paramiko.SSHClient()
    ssh_db1.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh_db1.connect(hostname='172.16.9.71', port=22, username='root', password='Wisedu#123')
    #db_name列表生成
    db_name_stdin, db_name_origin, db_name_stderr = ssh_db1.exec_command("source /home/oracle/.bash_profile;cat /etc/oratab|grep `echo $ORACLE_HOME`|grep -v Backup|awk -F ':' '{print $1}'")
    db_name = db_name_origin.read().split("\n")
    print db_name
    #ssh_db1连接关闭
    ssh_db1.close()
    #查询表空间使用率sql语句
    tablespace_sql = "select Upper(F.TABLESPACE_NAME),D.TOT_GROOTTE_MB,D.TOT_GROOTTE_MB - F.TOTAL_BYTES,To_char(Round(( D.TOT_GROOTTE_MB - F.TOTAL_BYTES ) / D.TOT_GROOTTE_MB *100, 2), '990.99')||'%',F.TOTAL_BYTES,F.MAX_BYTES FROM (SELECT TABLESPACE_NAME,Round(Sum(BYTES) / ( 1024*1024 ), 2) TOTAL_BYTES,Round(Max(BYTES) / ( 1024*1024 ), 2) MAX_BYTES FROM   SYS.DBA_FREE_SPACE GROUP  BY TABLESPACE_NAME) F,(SELECT DD.TABLESPACE_NAME,Round(Sum(DD.BYTES) / ( 1024*1024 ), 2) TOT_GROOTTE_MB FROM   SYS.DBA_DATA_FILES DD GROUP  BY DD.TABLESPACE_NAME) D WHERE  D.TABLESPACE_NAME = F.TABLESPACE_NAME ORDER  BY 1"
    #tablespace_sql = "select dbid,name,created,log_mode,open_mode from v$database"
#word路径，函数定义
    doc_path = r'check_database.docx'
    new_doc(userinfo,host,db_name,tablespace_sql,doc_path)