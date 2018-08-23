# -*- coding: utf-8 -*-
import os
from docx import  Document
from docx.shared import  Pt
from docx.oxml.ns import  qn
from docx.shared import Inches
import cx_Oracle
import paramiko
import sys
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import *
from docx.oxml.ns import qn
from docx.shared import RGBColor
import time
import datetime
import pytz
from profile import userinfo,host,school_name,author,frequency,ip1,ip2,port1,port2,username1,username2,password1,password2
#数据库连接函数
def connect_db(userinfo,host,service_name,sql):
    conn_str = userinfo + host + service_name
    conn = cx_Oracle.connect(conn_str)    
    c = conn.cursor()                                       
    c.execute(sql)
    if sql.startswith('select'):
        fir=c.fetchall()
        conn.commit()             
        c.close()                                 
        conn.close()
        return fir
    else:
        conn.commit()             
        c.close()                                 
        conn.close()
#输出asm磁盘组信息函数
def get_asm(userinfo,host,service_name,sql):
    document = Document(doc_path)
    asm = connect_db(userinfo,host,service_name[1],sql)
    asm_output = document.add_table(1,6,style="Medium Grid 3 Accent 1")
    asm_output_cells = asm_output.rows[0].cells
    asm_output_cells[0].text = u'Grp'
    asm_output_cells[1].text = u'Name'
    asm_output_cells[2].text = u'State'
    asm_output_cells[3].text = u'Type'
    asm_output_cells[4].text = u'Total(MB)'
    asm_output_cells[5].text = u'Free(MB)'
    for i in range(0,len(asm)):
        cells = asm_output.add_row().cells
        cells[0].text = str(asm[i][0])
        cells[1].text = str(asm[i][1])
        cells[2].text = str(asm[i][2])
        cells[3].text = str(asm[i][3])
        cells[4].text = str(asm[i][4])
        cells[5].text = str(asm[i][5])
    document.save(doc_path)
#输出数据库版本函数
def get_database_version(userinfo,host,service_name,sql):
    document = Document(doc_path)
    database_version = connect_db(userinfo,host,service_name[1],sql)
    database_version_output = document.add_table(1,1,style="Medium Grid 3 Accent 1")
    database_version_output_cells = database_version_output.rows[0].cells
    database_version_output_cells[0].text = u'BANNER'
    for i in range(0,len(database_version)):
        cells = database_version_output.add_row().cells
        cells[0].text = str(database_version[i][0])
    document.save(doc_path)
#输出SGA组成函数
def get_sga(userinfo,host,service_name,sql):
    document = Document(doc_path)
    names = locals()
    for q in service_name:
        if q.strip():
            sga = connect_db(userinfo,host,q,sql)
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s数据库SGA组成：" % q)
            r_total.font.bold = True
            names["sga_output%s" % q]= document.add_table(1,2,style="Medium Grid 3 Accent 1")
            sga_output_cells = names["sga_output%s" % q].rows[0].cells
            sga_output_cells[0].text = u'NAME'
            sga_output_cells[1].text = u'VALUE'
            for i in range(0,len(sga)):
                cells = names["sga_output%s" % q].add_row().cells
                cells[0].text = str(sga[i][0])
                cells[1].text = str(sga[i][1])
        else:
            pass
    document.save(doc_path)
#输出控制文件信息函数
def get_controlfile(userinfo,host,service_name,sql):
    document = Document(doc_path)
    names = locals()
    for q in service_name:
        if q.strip():
            controlfile = connect_db(userinfo,host,q,sql)
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s数据库控制文件：" % q)
            r_total.font.bold = True
            names["controlfile_output%s" % q]= document.add_table(1,1,style="Medium Grid 3 Accent 1")
            controlfile_output_cells = names["controlfile_output%s" % q].rows[0].cells
            controlfile_output_cells[0].text = u'NAME'
            for i in range(0,len(controlfile)):
                cells = names["controlfile_output%s" % q].add_row().cells
                cells[0].text = str(controlfile[i][0])
        else:
            pass
    document.save(doc_path)
#输出在线重做日志信息函数
def get_redolog(userinfo,host,service_name,sql):
    document = Document(doc_path)
    names = locals()
    for q in service_name:
        if q.strip():
            redolog = connect_db(userinfo,host,q,sql)
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s数据库在线重做日志：" % q)
            r_total.font.bold = True
            names["redolog_output%s" % q]= document.add_table(1,4,style="Medium Grid 3 Accent 1")
            redolog_output_cells = names["redolog_output%s" % q].rows[0].cells
            redolog_output_cells[0].text = u'Group#'
            redolog_output_cells[1].text = u'Status'
            redolog_output_cells[2].text = u'Type'
            redolog_output_cells[3].text = u'Member'
            for i in range(0,len(redolog)):
                cells = names["redolog_output%s" % q].add_row().cells
                cells[0].text = str(redolog[i][0])
                cells[1].text = str(redolog[i][1])
                cells[2].text = str(redolog[i][2])
                cells[3].text = str(redolog[i][3])
        else:
            pass
    document.save(doc_path)
#输出表空间使用率函数
def get_tablespace(userinfo,host,service_name,sql):
    document = Document(doc_path)
    names = locals()
    for q in service_name:
        if q.strip():
            tablespace = connect_db(userinfo,host,q,sql)
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s数据库表空间使用率：" % q)
            r_total.font.bold = True
            names["tablespace_output%s" % q]= document.add_table(1,6,style="Medium Grid 3 Accent 1")
            tablespace_output_cells = names["tablespace_output%s" % q].rows[0].cells
            tablespace_output_cells[0].text = u'表空间名'
            tablespace_output_cells[1].text = u'总空间'
            tablespace_output_cells[2].text = u'已使用'
            tablespace_output_cells[3].text = u'使用率'
            tablespace_output_cells[4].text = u'剩余空间'
            tablespace_output_cells[5].text = u'最大块'
            for i in range(0,len(tablespace)):
                cells = names["tablespace_output%s" % q].add_row().cells
                cells[0].text = str(tablespace[i][0])
                cells[1].text = str(tablespace[i][1])
                cells[2].text = str(tablespace[i][2])
                cells[3].text = str(tablespace[i][3])
                cells[4].text = str(tablespace[i][4])
                cells[5].text = str(tablespace[i][5])
        else:
            pass
    document.save(doc_path)
#输出病毒检测结果函数
def get_virus(userinfo,host,service_name,sql):
    document = Document(doc_path)
    names = locals()
    for q in service_name:
        if q.strip():
            virus = connect_db(userinfo,host,q,sql)
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s数据库已知病毒检测结果：" % q)
            r_total.font.bold = True
            if virus:
                document.add_paragraph(u'%s数据库受到病毒感染。' % q)
                names["virus_output%s" % q]= document.add_table(1,3,style="Medium Grid 3 Accent 1")
                virus_output_cells = names["virus_output%s" % q].rows[0].cells
                virus_output_cells[0].text = u'病毒所属用户'
                virus_output_cells[1].text = u'病毒对象名'
                virus_output_cells[2].text = u'病毒对象类型'
                for i in range(0,len(virus)):
                    cells = names["virus_output%s" % q].add_row().cells
                    cells[0].text = str(virus[i][0])
                    cells[1].text = str(virus[i][1])
                    cells[2].text = str(virus[i][2])
            else:
                document.add_paragraph(u'%s数据库未受到已知病毒感染。' % q)
        else:
            pass
    document.save(doc_path)
#输出节点一数据库alert日志
def get_database_alert1(oracle_base,service_name,hostname,port,username,password):
    document = Document(doc_path)
    names = locals()
    ssh_db1 = paramiko.SSHClient()
    ssh_db1.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh_db1.connect(hostname, port, username, password)
    for q in service_name:
        if q.strip():
            cmd_database_alert1_1 = "tail -1000 %s" % oracle_base
            cmd_database_alert1_2 = "/diag/rdbms/%s/" % q
            cmd_database_alert1_3 = "%s1/trace/" % q
            cmd_database_alert1_4 = "alert_%s1.log|grep ORA-" % q
            cmd_database_alert1 = cmd_database_alert1_1 + cmd_database_alert1_2 + cmd_database_alert1_3 + cmd_database_alert1_4
            print cmd_database_alert1
            database_alert1_stdin, names["database_alert1_origin_%s" % q], database_alert1_stderr = ssh_db1.exec_command(cmd_database_alert1)
            names["database_alert1_%s" % q] = names["database_alert1_origin_%s" % q].read()
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s1实例alert告警日志：" % q)
            r_total.font.bold = True
            p_total = document.add_paragraph()
            p_total = p_total.add_run(names["database_alert1_%s" % q])
    document.save(doc_path)
#输出节点二数据库alert日志
def get_database_alert2(oracle_base,service_name,hostname,port,username,password):
    document = Document(doc_path)
    names = locals()
    ssh_db2 = paramiko.SSHClient()
    ssh_db2.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh_db2.connect(hostname, port, username, password)
    for q in service_name:
        if q.strip():
            cmd_database_alert2_1 = "tail -1000 %s" % oracle_base
            cmd_database_alert2_2 = "/diag/rdbms/%s/" % q
            cmd_database_alert2_3 = "%s2/trace/" % q
            cmd_database_alert2_4 = "alert_%s2.log|grep ORA-" % q
            cmd_database_alert2 = cmd_database_alert2_1 + cmd_database_alert2_2 + cmd_database_alert2_3 + cmd_database_alert2_4
            database_alert2_stdin, names["database_alert2_origin_%s" % q], database_alert2_stderr = ssh_db2.exec_command(cmd_database_alert2)
            print cmd_database_alert2
            names["database_alert2_%s" % q] = names["database_alert2_origin_%s" % q].read()
            p_total = document.add_paragraph()
            r_total = p_total.add_run(u"%s2实例alert告警日志：" % q)
            r_total.font.bold = True
            p_total = document.add_paragraph()
            p_total = p_total.add_run(names["database_alert2_%s" % q])
    document.save(doc_path)
#word输出结构	
def new_doc(userinfo,host,school_name,author,frequency,ip1,ip2,port1,port2,username1,username2,password1,password2,time,db_name,db_name_output,oracle_base,cpu1,cpu2,memory1,memory2,swap1,swap2,gateway1,gateway2,version,os_version,timezone1,timezone2,syslog1,syslog2,filesystem1,filesystem2,cluster,cluster_alert1,cluster_alert2,crsdlog1,crsdlog2,cssdlog1,cssdlog2,ocr_backup,asm_sql,asm_alert1,asm_alert2,database_version_sql,sga_sql,controlfile_sql,redolog_sql,tablespace_sql,virus_sql,listener1,listener2,listener_log1,listener_log2,doc_path):#新建一个word文档,写入汇总表的数据
    document = Document()
    document.styles['Title'].font.name = u'Microsoft YaHei'
    document.styles['Title'].font.size = Pt(26)
    document.styles['Title'].font.underline = False
    document.styles['Title']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft YaHei')
    document.styles['Title'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.styles['Normal'].font.name = u'等线 Light'
    document.styles['Normal'].font.size = Pt(9)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线 Light')
    document.styles['Normal'].paragraph_format.left_indent = Inches(0)
    #document.styles['Normal'].paragraph_format.first_line_indent = Inches(0.25)
    #document.styles['Normal'].paragraph_format.space_before = Pt(8)
    #document.styles['Normal'].paragraph_format.space_after = Pt(8)
    #document.styles['Normal'].paragraph_format.line_spacing = 1
    document.styles['Heading 1'].font.name = u'等线 Light'
    document.styles['Heading 1'].font.size = Pt(22)
    document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线 Light')
    document.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    #document.styles['Heading 1'].next_paragraph_style = document.styles['Normal']
    document.styles['Heading 2'].font.name = u'等线 Light'
    document.styles['Heading 2'].font.size = Pt(16)
    document.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线 Light')
    document.styles['Heading 3'].font.name = u'等线 Light'
    document.styles['Heading 3'].font.size = Pt(14)
    document.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.styles['Heading 3']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线 Light')
    #document.add_heading(u'1 巡检说明',level=1)
    document.add_paragraph(u'\n\n\n\n%s\n\nOracle数据库\n\n巡\n\n检\n\n报\n\n告' % school_name,style='Title').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(u'        版权声明和保密须知\n        本文中出现的任何文字描述、文档格式、插图、照片、方法、过程等内容，除另有特别注明，版权均属江苏金智教育信息股份有限公司所有，受到有关产权及版权法保护。任何单位和个人未经江苏金智教育信息股份有限公司的书面授权许可，不得复制或引用本文件的任何片段，无论通过电子形式或非电子形式。\n        Copyright @ 2018 江苏金智教育信息股份有限公司 版权所有',style='Normal')
    document.add_page_break()
    document.add_heading(u'文档控制',level=1)
    document.add_paragraph(u'        此文档仅供%s与江苏金智教育信息股份有限公司审阅，不得向与此无关的个人或机构传阅或复制。' % school_name,style='Normal')
    document.add_heading(u'修改记录',level=1)
    record_output = document.add_table(1,4,style="Table Grid")
    record_output_cells = record_output.rows[0].cells
    record_output_cells[0].text = u'日期'
    record_output_cells[1].text = u'作者'
    record_output_cells[2].text = u'版本'
    record_output_cells[3].text = u'修改记录'
    cells = record_output.add_row().cells
    cells[0].text = time
    cells[1].text = author
    cells[2].text = u'V1.0'
    cells[3].text = u'起草'
    cells = record_output.add_row().cells
    document.add_heading(u'审阅记录',level=1)
    review_output = document.add_table(1,2,style="Table Grid")
    review_output_cells = review_output.rows[0].cells
    review_output_cells[0].text = u'姓名'
    review_output_cells[1].text = u'职务'
    cells = review_output.add_row().cells
    cells[0].text = u'胡楠'
    cells[1].text = u'数据库服务部经理'
    document.add_heading(u'相关文档',level=1)
    document.add_paragraph(u'        无')
    document.add_page_break()
    document.add_paragraph(u'        添加目录')
    document.add_page_break()
    document.add_heading(u'1 巡检说明',level=1)
    document.add_heading(u'1.1 基本信息',level=2)
    document.add_paragraph(u'本次总共对%s' % school_name + u'%s套Oracle数据库进行巡检。' % (len(db_name)-2),style='Normal')
    document.add_paragraph(u'巡检时间：' + str(time))
    document.add_paragraph(u'巡检方式：■远程巡检  □现场巡检。')
    document.add_paragraph(u'巡检的数据库列表如下：')
    info_output = document.add_table(1,4,style="Medium Grid 3 Accent 1")
    info_output_cells = info_output.rows[0].cells
    info_output_cells[0].text = u'数据库名及版本'
    info_output_cells[1].text = u'IP地址'
    info_output_cells[2].text = u'操作系统'
    info_output_cells[3].text = u'描述'
    for q in db_name:
        if q.strip():
            cells = info_output.add_row().cells
            cells[0].text = str(q) + '\n' + 'Oracle' + ' ' + str(version).replace('\n','')
            cells[1].text = str(ip1) + '\n' + str(ip2)
            cells[2].text = str(os_version)
            cells[3].text = ''
        else:
            pass
    document.add_heading(u'1.2 巡检小结',level=2)
    summary_output = document.add_table(1,3,style="Medium Grid 3 Accent 1")
    summary_output_cells = summary_output.rows[0].cells
    summary_output_cells[0].text = u'类别'
    summary_output_cells[1].text = u'检查细项'
    summary_output_cells[2].text = u'结果'
    summary_output_cells[0].width = Inches(0.2)
    summary_output_cells[1].width = Inches(1.0)
    summary_output_cells[2].width = Inches(4.0)
    cells = summary_output.add_row().cells
    cells[0].text = u'操作系统巡检'
    cells[1].text = u'操作系统基本信息'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'操作系统时区'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'操作系统日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'操作系统磁盘空间'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u'集群巡检'
    cells[1].text = u'集群资源状态'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'集群告警日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'集群CRS日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'集群CSS日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'OCR自动备份'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u'ASM巡检'
    cells[1].text = u'磁盘组信息'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'ASM告警日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u'数据库巡检'
    cells[1].text = u'数据库版本'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'SGA组成'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'控制文件'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'在线重做日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'表空间管理'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'已知木马检测'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'数据库告警日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u'网络巡检'
    cells[1].text = u'监听状态'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'监听日志'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u'备份策略评估'
    cells[1].text = u'RMAN物理备份'
    cells[2].text = u'■通过  □警告  □失败'
    cells = summary_output.add_row().cells
    cells[0].text = u''
    cells[1].text = u'逻辑备份'
    cells[2].text = u'■通过  □警告  □失败'
    document.add_heading(u'1.3 问题点与调整',level=2)
    document.add_paragraph(u'        根据本次巡检结果，目前Oracle数据库环境存在以下几个问题点：',style='Normal')
    document.add_heading(u'1.4 后续建议',level=2)
    document.add_paragraph(u'        针对巡检结果，后续建议：',style='Normal')
    document.add_heading(u'2 巡检细项',level=1)
    document.add_heading(u'2.1 操作系统巡检',level=2)
    document.add_heading(u'2.1.1 操作系统基本信息',level=3)
    sysinfo_output = document.add_table(1,5,style="Medium Grid 3 Accent 1")
    sysinfo_output_cells = sysinfo_output.rows[0].cells
    sysinfo_output_cells[0].text = u'主机名'
    sysinfo_output_cells[1].text = u'CPU配置'
    sysinfo_output_cells[2].text = u'物理内存'
    sysinfo_output_cells[3].text = u'交换分区'
    sysinfo_output_cells[4].text = u'默认网关'
    cells = sysinfo_output.add_row().cells
    cells[0].text = str(hostname1)
    cells[1].text = str(cpu1) + ' ' + 'Processor(s)'
    cells[2].text = str(memory1)
    cells[3].text = str(swap1)
    cells[4].text = str(gateway1)
    cells = sysinfo_output.add_row().cells
    cells[0].text = str(hostname2)
    cells[1].text = str(cpu2) + ' ' + 'Processor(s)'
    cells[2].text = str(memory2)
    cells[3].text = str(swap2)
    cells[4].text = str(gateway2)
#输出节点一时区
    document.add_heading(u'2.1.2 操作系统时间与时区',level=3)
    document.add_paragraph(u'        数据库服务器的操作系统时区与时间设置，要符合本地时区时间设置。否则对部分时间敏感业务会造成一定的影响。譬如，教务选课系统、一卡通系统等。',style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一时间与时区：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(timezone1)
#输出节点二时区
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二时间与时区：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(timezone2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出节点一操作系统文件系统使用率
    document.add_heading(u'2.1.3 操作系统磁盘空间',level=3)
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一文件系统使用率：')
    r_total.font.bold = True
    filesystem1_output = document.add_table(1,7,style="Medium Grid 3 Accent 1")
    filesystem1_output_cells = filesystem1_output.rows[0].cells
    filesystem1_output_cells[0].text = u'Filesystem'
    filesystem1_output_cells[1].text = u'GB blocks'
    filesystem1_output_cells[2].text = u'Free'
    filesystem1_output_cells[3].text = u'%Used'
    filesystem1_output_cells[4].text = u'Iused'
    filesystem1_output_cells[5].text = u'%Iused'
    filesystem1_output_cells[6].text = u'Mounted on'
    for i in range(1,len(filesystem1)/7):
        cells = filesystem1_output.add_row().cells
        cells[0].text = str(filesystem1[i*7+2])
        cells[1].text = str(filesystem1[i*7+3])
        cells[2].text = str(filesystem1[i*7+4])
        cells[3].text = str(filesystem1[i*7+5])
        cells[4].text = str(filesystem1[i*7+6])
        cells[5].text = str(filesystem1[i*7+7])
        cells[6].text = str(filesystem1[i*7+8])
#输出节点二操作系统文件系统使用率
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二文件系统使用率：')
    r_total.font.bold = True
    filesystem2_output = document.add_table(1,7,style="Medium Grid 3 Accent 1")
    filesystem2_output_cells = filesystem2_output.rows[0].cells
    filesystem2_output_cells[0].text = u'Filesystem'
    filesystem2_output_cells[1].text = u'GB blocks'
    filesystem2_output_cells[2].text = u'Free'
    filesystem2_output_cells[3].text = u'%Used'
    filesystem2_output_cells[4].text = u'Iused'
    filesystem2_output_cells[5].text = u'%Iused'
    filesystem2_output_cells[6].text = u'Mounted on'
    for i in range(1,len(filesystem2)/7):
        cells = filesystem2_output.add_row().cells
        cells[0].text = str(filesystem2[i*7+2])
        cells[1].text = str(filesystem2[i*7+3])
        cells[2].text = str(filesystem2[i*7+4])
        cells[3].text = str(filesystem2[i*7+5])
        cells[4].text = str(filesystem2[i*7+6])
        cells[5].text = str(filesystem2[i*7+7])
        cells[6].text = str(filesystem2[i*7+8])
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出节点一操作系统日志
    document.add_heading(u'2.1.4 操作系统日志',level=3)
    document.add_paragraph(u'        操作系统日志记录了操作系统运行相关的日志信息。当硬件发生变更、软件及操作系统服务运行异常时，系统日志均会予以记录。本次巡检，对最近%s个月的操作系统运行日志予以检查。' % frequency,style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一操作系统日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(syslog1)
#输出节点二操作系统日志
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二操作系统日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(syslog2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出集群资源状态
    document.add_heading(u'2.2 集群巡检',level=2)
    document.add_heading(u'2.2.1 集群资源状态',level=3)
    p_total = document.add_paragraph()
    p_total = p_total.add_run(cluster)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出节点一集群告警日志
    document.add_heading(u'2.2.2 集群告警日志',level=3)
    document.add_paragraph(u'        集群告警日志包含了集群系统级别的错误与报警，需要对近期发现的集群错误进行检查与分析。本次巡检时，在对集群两个节点的alert日志进行分析，发现最近%s个月，集群告警日志无异常记录。' % frequency,style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一集群告警日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(cluster_alert1)
#输出节点二集群告警日志
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二集群告警日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(cluster_alert2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出节点一集群crs日志
    document.add_heading(u'2.2.3 集群CRS日志',level=3)
    document.add_paragraph(u'        集群CRS日志包含了集群CRS资源的错误与告警，需要对近期发现的集群资源错误进行检查与分析。本次巡检，对集群两个节点的CRS日志进行分析，发现最近%s个月，集群CRS日志无异常记录。' % frequency,style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一crs日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(crsdlog1)
#输出节点二集群crs日志
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二crs日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(crsdlog2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出节点一集群css日志
    document.add_heading(u'2.2.4 集群CSS日志',level=3)
    document.add_paragraph(u'        集群CSS日志包含了集群节点通信、同步的错误与告警，需要对近期发现的集群节点通信、同步错误进行检查与分析。本次巡检时，在对集群两个节点的CSS日志进行分析，发现最近%s个月，集群CSS日志无异常记录。' % frequency,style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一css日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(cssdlog1)
#输出节点二集群css日志
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二css日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(cssdlog2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出ocr备份状态
    document.add_heading(u'2.2.5 OCR自动备份状态',level=3)
    document.add_paragraph(u'        巡检过程中，对OCR自动备份进行检查。本次巡检，OCR自动备份正常',style='Normal')
    p_total = document.add_paragraph()
    p_total = p_total.add_run(ocr_backup)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出ASM磁盘组信息
    document = Document(doc_path)
    document.add_heading(u'2.3 ASM巡检',level=2)
    document.add_paragraph(u'        本节主要对ASM实例运行状态、磁盘组的配置情况进行详细描述与检查。',style='Normal')
    document.add_heading(u'2.3.1 磁盘组信息',level=3)
    document.save(doc_path)
    get_asm(userinfo,host,db_name,asm_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出节点一ASM ALERT告警日志
    document = Document(doc_path)
    document.add_heading(u'2.3.2 ASM告警日志',level=3)
    document.add_paragraph(u'        检查了集群ASM实例最近%s个月的alert告警日志，告警日志无异常记录。' % frequency,style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一ASM ALERT告警日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(asm_alert1)
#输出节点二ASM ALERT告警日志
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二ASM ALERT告警日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(asm_alert2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出数据库版本
    document = Document(doc_path)
    document.add_heading(u'2.4 数据库巡检',level=2)
    document.add_paragraph(u'        该部分详细阐述了数据库%s的主要结构。' % db_name_output + u'注：以下部分，若没有特殊说明，均表示%s个数据库配置一致。' % (len(db_name)-1),style='Normal')
    document.add_heading(u'2.4.1 数据库版本',level=3)
    document.save(doc_path)
    get_database_version(userinfo,host,db_name,database_version_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出sga
    document = Document(doc_path)
    document.add_heading(u'2.4.2 SGA组成',level=3)
    document.add_paragraph(u'        以下是数据库%s的SGA内存参数设置信息：' % db_name_output,style='Normal')
    document.save(doc_path)
    get_sga(userinfo,host,db_name,sga_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出控制文件信息
    document = Document(doc_path)
    document.add_heading(u'2.4.3 控制文件',level=3)
    document.add_paragraph(u'        以下是数据库%s的控制文件相关信息：' % db_name_output,style='Normal')
    document.save(doc_path)
    get_controlfile(userinfo,host,db_name,controlfile_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出在线重做日志信息
    document = Document(doc_path)
    document.add_heading(u'2.4.4 在线重做日志',level=3)
    document.add_paragraph(u'        以下是数据库%s的在线重做日志相关信息：' % db_name_output,style='Normal')
    document.save(doc_path)
    get_redolog(userinfo,host,db_name,redolog_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出表空间使用率
    document = Document(doc_path)
    document.add_heading(u'2.4.5 表空间管理',level=3)
    document.add_paragraph(u'        以下是数据库表空间的管理和使用情况。在ORACLE 9i之后除了系统表空间，其他的表空间的空间段管理方式推荐为LOCAL。\n        临时表空间用于存放临时段。为了维护数据库的性能，临时表空间的维护方法有别于其他一般表空间。缺省情况下，所有表空间都创建为PERMANENT。所以在创建临时段时，需要保证表空间类型为TEMPORARY。由于这些表空间中的排序段不被清除，所以减少了空间事务争夺，同时减少了SMON对于CPU的使用率。\n        由于表空间的extent 出现了local management 方式，对表空间采用了位图管理，更利于空间的使用及回收管理。',style='Normal')
    document.save(doc_path)
    get_tablespace(userinfo,host,db_name,tablespace_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出病毒检查结果
    document = Document(doc_path)
    document.add_heading(u'2.4.6 已知木马检测',level=3)
    document.add_paragraph(u'        根据各大安全网站公布的数据，定期更新数据库病毒检测库，对已知的数据库病毒程序进行检查，本次巡检数据库安全情况如下：',style='Normal')
    document.save(doc_path)
    get_virus(userinfo,host,db_name,virus_sql)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出节点一数据库alert告警日志
    document = Document(doc_path)
    document.add_heading(u'2.4.7 数据库告警日志',level=3)
    document.add_paragraph(u'        数据库alert告警日志包含了系统级别的错误与告警，需要对发现的ORA错误进行检查与分析。最近%s个月告警日志无异常记录。' % frequency,style='Normal')
    document.save(doc_path)
    get_database_alert1(oracle_base,db_name,ip1,port1,username1,password1)
#输出节点二数据库alert告警日志
    get_database_alert2(oracle_base,db_name,ip2,port2,username2,password2)
    document = Document(doc_path)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
    document.save(doc_path)
#输出节点一监听状态
    document = Document(doc_path)
    document.add_heading(u'2.5 网络巡检',level=2)
    document.add_paragraph(u'        本节主要对数据库网络、监听的运行情况予以检查。',style='Normal')
    document.add_heading(u'2.5.1 监听状态',level=3)
    document.add_paragraph(u'        巡检过程中，对监听状态进行检查，结果显示，两个节点的监听状态及服务均正常。',style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一监听状态：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(listener1)
#输出节点二监听状态
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二监听状态：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(listener2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#输出节点一监听日志
    document.add_heading(u'2.5.2 监听日志',level=3)
    document.add_paragraph(u'        巡检过程中，对集群两个节点的监听日志予以检查。检查发现，最近%s个月两个节点的监听日志无异常记录。' % frequency,style='Normal')
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点一监听日志：')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(listener_log1)
#输出节点二监听日志
    p_total = document.add_paragraph()
    r_total = p_total.add_run(u'节点二监听日志')
    r_total.font.bold = True
    p_total = document.add_paragraph()
    p_total = p_total.add_run(listener_log2)
    document.add_paragraph(u'结论：■正常  □不正常。',style='Normal')
#保存word
    document.add_heading(u'2.6 数据库备份策略评估',level=2)
    document.add_paragraph(u'        本节主要对数据库的备份与恢复策略进行检查与评估。',style='Normal')
    backup_output = document.add_table(1,8,style="Medium Grid 3 Accent 1")
    backup_output.alignment = WD_TABLE_ALIGNMENT.LEFT
    backup_output.allow_autofit = False
    backup_output_cells = backup_output.rows[0].cells
    backup_output_cells[0].text = u'数据库用途'
    backup_output_cells[1].text = u'数据库名'
    backup_output_cells[2].text = u'物理备份情况（含异地备份）'
    backup_output_cells[3].text = u'物理备份保留周期'
    backup_output_cells[4].text = u'物理备份数据可恢复时间段'
    backup_output_cells[5].text = u'逻辑备份情况（含异地备份）'
    backup_output_cells[6].text = u'逻辑备份保留周期'
    backup_output_cells[7].text = u'逻辑备份数据可恢复时间段'
    backup_output_cells[0].width = Inches(1.0)
    for i in range(0,(len(db_name)-1)):
        cells = backup_output.add_row().cells
        cells[0].text = ''
        cells[1].text = str(db_name[i])
        cells[2].text = u'■正常\n□警告\n□异常'
        cells[3].text = u'本地备份保留14天，异地备份保留14天'
        cells[4].text = u'可以恢复到14天之内任意时间点'
        cells[5].text = u'■正常\n□警告\n□异常'
        cells[6].text = u'本地备份保留7天，异地备份保留7天'
        cells[7].text = u'可以恢复到7天内做逻辑备份的时间点'
    document.save(doc_path)
#主函数
if __name__ == "__main__":
#sql语句
	#查询asm磁盘组信息sql语句
    asm_sql = "select GROUP_NUMBER,name,state,type,total_MB,free_MB from v$asm_diskgroup"
    #asm = get_data(asm_sql)
    #查询数据库版本sql语句
    database_version_sql = "select * from v$version"
    #database_version = get_data(database_version_sql)
    #查询sga sql语句
    sga_sql = "select * from v$sga"
    #查询控制文件sql
    controlfile_sql = "select name from v$controlfile"
    #查询redolog sql
    redolog_sql = "select group#,status,type,member from v$logfile"
    #查询表空间使用率sql语句
    tablespace_sql = "select Upper(F.TABLESPACE_NAME),D.TOT_GROOTTE_MB,D.TOT_GROOTTE_MB - F.TOTAL_BYTES,To_char(Round(( D.TOT_GROOTTE_MB - F.TOTAL_BYTES ) / D.TOT_GROOTTE_MB *100, 2), '990.99')||'%',F.TOTAL_BYTES,F.MAX_BYTES FROM (SELECT TABLESPACE_NAME,Round(Sum(BYTES) / ( 1024*1024 ), 2) TOTAL_BYTES,Round(Max(BYTES) / ( 1024*1024 ), 2) MAX_BYTES FROM   SYS.DBA_FREE_SPACE GROUP  BY TABLESPACE_NAME) F,(SELECT DD.TABLESPACE_NAME,Round(Sum(DD.BYTES) / ( 1024*1024 ), 2) TOT_GROOTTE_MB FROM   SYS.DBA_DATA_FILES DD GROUP  BY DD.TABLESPACE_NAME) D WHERE  D.TABLESPACE_NAME = F.TABLESPACE_NAME ORDER  BY 1"
    #查询已知木马sql
    virus_sql = "select owner,object_name,object_type from dba_objects where object_name like '%DBMS_SUPPORT_INTERNAL%' or object_name like '%DBMS_SYSTEM_INTERNAL%' or object_name like '%DBMS_STANDARD_FUN9%' or object_name like '%DBMS_CORE_INTERNAL%' or object_name like '%DBMS_SUPPORT_DBMONITOR%' order by object_name"
    #北京时间
    tz = pytz.timezone('Asia/Shanghai') #东八区
    time = datetime.datetime.fromtimestamp(int(time.time()), pytz.timezone('Asia/Shanghai')).strftime('%Y-%m-%d %H:%M:%S')
    print(isinstance(time,str))
#ssh连接到节点一取的信息
    ssh_db1 = paramiko.SSHClient()
    ssh_db1.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh_db1.connect(ip1, port1, username1, password1)
    #hostname1变量生成
    hostname1_stdin, hostname1_origin, hostname1_stderr = ssh_db1.exec_command("hostname")
    hostname1 = hostname1_origin.read().replace('\n','')
    #cpu1变量生成
    cpu1_stdin, cpu1_origin, cpu1_stderr = ssh_db1.exec_command("prtconf|grep -i 'Number Of Processors'|awk -F ':' '{print $2}'")
    cpu1 = cpu1_origin.read().replace('\n','')
    #memory1变量生成
    memory1_stdin, memory1_origin, memory1_stderr = ssh_db1.exec_command("prtconf|grep -i 'Good Memory Size'|awk -F ':' '{print $2}'")
    memory1 = memory1_origin.read().replace('\n','')
    #swap1变量生成
    swap1_stdin, swap1_origin, swap1_stderr = ssh_db1.exec_command("prtconf|grep -i 'Total Paging Space'|awk -F ':' '{print $2}'")
    swap1 = swap1_origin.read().replace('\n','')
    #gateway1变量生成
    gateway1_stdin, gateway1_origin, gateway1_stderr = ssh_db1.exec_command("netstat -rn|grep default|awk -F ' ' '{print $2}'")
    gateway1 = gateway1_origin.read().replace('\n','')
    #db_name列表生成
    db_name_stdin, db_name_origin, db_name_stderr = ssh_db1.exec_command("source /home/oracle/.profile;cat /etc/oratab|grep `echo $ORACLE_HOME`|grep -v Backup|awk -F ':' '{print $1}'")
    db_name = db_name_origin.read().split("\n")
    print db_name
    db_name_output = ''
    for d in db_name:
        if d.strip():
            db_name_output = db_name_output + ' ' + d
        else:
            pass
    print db_name_output
    #oracle_base生成
    oracle_base_stdin, oracle_base_origin, oracle_base_stderr = ssh_db1.exec_command("source /home/oracle/.profile;echo $ORACLE_BASE")
    oracle_base = oracle_base_origin.read().replace('\n','')
    #oracle_home生成
    oracle_home_stdin, oracle_home_origin, oracle_home_stderr = ssh_db1.exec_command("source /home/oracle/.profile;echo $ORACLE_HOME")
    oracle_home = oracle_home_origin.read().replace('\n','')
    #grid_base生成
    grid_base_stdin, grid_base_origin, grid_base_stderr = ssh_db1.exec_command("source /home/grid/.profile;echo $ORACLE_BASE")
    grid_base = grid_base_origin.read().replace('\n','')
    #grid_home生成
    grid_home_stdin, grid_home_origin, grid_home_stderr = ssh_db1.exec_command("source /home/grid/.profile;echo $ORACLE_HOME")
    grid_home = grid_home_origin.read().replace('\n','')
    #timezone1变量生成
    timezone1_stdin, timezone1_origin, timezone1_stderr = ssh_db1.exec_command("date '+%Y-%m-%d %H:%M:%S';echo $TZ")
    timezone1 = timezone1_origin.read()
    #syslog1变量生成
    syslog1_stdin, syslog1_origin, syslog1_stderr = ssh_db1.exec_command("errpt")
    syslog1 = syslog1_origin.read()
    #filesystem1变量生成
    filesystem1_stdin, filesystem1_origin, filesystem1_stderr = ssh_db1.exec_command("df -g")
    filesystem1 = filesystem1_origin.read().split()
    #cluster变量生成
    cluster_stdin, cluster_origin, cluster_stderr = ssh_db1.exec_command("source /home/grid/.profile;crsctl stat res -t")
    cluster = cluster_origin.read()
    #cluster_alert1变量生成
    cmd_cluster_alert1_1 = "tail -200 %s/log/" % grid_home
    cmd_cluster_alert1_2 = "%s/" % hostname1
    cmd_cluster_alert1_3 = "alert%s.log" % hostname1
    cmd_cluster_alert1 = cmd_cluster_alert1_1 + cmd_cluster_alert1_2 + cmd_cluster_alert1_3
    print cmd_cluster_alert1
    cluster_alert1_stdin, cluster_alert1_origin, cluster_alert1_stderr = ssh_db1.exec_command(cmd_cluster_alert1)
    cluster_alert1 = cluster_alert1_origin.read()
    #crsdlog1变量生成
    cmd_crsdlog1_1 = "tail -200 %s" % grid_home
    cmd_crsdlog1_2 = "/log/%s/crsd/crsd.log" % hostname1
    cmd_crsdlog1 = cmd_crsdlog1_1 + cmd_crsdlog1_2
    print cmd_crsdlog1
    crsdlog1_stdin, crsdlog1_origin, crsdlog1_stderr = ssh_db1.exec_command(cmd_crsdlog1)
    crsdlog1 = crsdlog1_origin.read()
    #cssdlog1变量生成
    cmd_cssdlog1_1 = "tail -200 %s" % grid_home
    cmd_cssdlog1_2 = "/log/%s/cssd/ocssd.log" % hostname1
    cmd_cssdlog1 = cmd_cssdlog1_1 + cmd_cssdlog1_2
    print cmd_cssdlog1
    cssdlog1_stdin, cssdlog1_origin, cssdlog1_stderr = ssh_db1.exec_command(cmd_cssdlog1)
    cssdlog1 = cssdlog1_origin.read()
    #ocr_backup变量生成
    ocr_backup_stdin, ocr_backup_origin, ocr_backup_stderr = ssh_db1.exec_command("source /home/grid/.profile;ocrconfig -showbackup")
    ocr_backup = ocr_backup_origin.read()
    #asm_alert1变量生成
    cmd_asm_alert1 = "tail -200 %s/diag/asm/+asm/+ASM1/trace/alert_+ASM1.log|grep ORA-" % grid_base
    asm_alert1_stdin, asm_alert1_origin, asm_alert1_stderr = ssh_db1.exec_command(cmd_asm_alert1)
    asm_alert1 = asm_alert1_origin.read()
    #listener1变量生成
    listener1_stdin, listener1_origin, listener1_stderr = ssh_db1.exec_command("source /home/grid/.profile;lsnrctl stat")
    listener1 = listener1_origin.read()
    #listener_log1变量生成
    cmd_listener_log1_1 = "tail -1000 %s/diag/tnslsnr/" % grid_base
    cmd_listener_log1_2 = "%s/listener/trace/listener.log|grep TNS-" % hostname1
    cmd_listener_log1 = cmd_listener_log1_1 + cmd_listener_log1_2
    print cmd_listener_log1
    listener_log1_stdin, listener_log1_origin, listener_log1_stderr = ssh_db1.exec_command(cmd_listener_log1)
    listener_log1 = listener_log1_origin.read()
    #ssh_db1连接关闭
    ssh_db1.close()
#ssh连接到节点二取的信息    
    ssh_db2 = paramiko.SSHClient()
    ssh_db2.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh_db2.connect(ip2, port2, username2, password2)
    #hostname2变量生成
    hostname2_stdin, hostname2_origin, hostname2_stderr = ssh_db2.exec_command("hostname")
    hostname2 = hostname2_origin.read().replace('\n','')
    #cpu2变量生成
    cpu2_stdin, cpu2_origin, cpu2_stderr = ssh_db2.exec_command("prtconf|grep -i 'Number Of Processors'|awk -F ':' '{print $2}'")
    cpu2 = cpu2_origin.read().replace('\n','')
    #memory2变量生成
    memory2_stdin, memory2_origin, memory2_stderr = ssh_db2.exec_command("prtconf|grep -i 'Good Memory Size'|awk -F ':' '{print $2}'")
    memory2 = memory2_origin.read().replace('\n','')
    #swap2变量生成
    swap2_stdin, swap2_origin, swap2_stderr = ssh_db2.exec_command("prtconf|grep -i 'Total Paging Space'|awk -F ':' '{print $2}'")
    swap2 = swap2_origin.read().replace('\n','')
    #gateway2变量生成
    gateway2_stdin, gateway2_origin, gateway2_stderr = ssh_db2.exec_command("netstat -rn|grep default|awk -F ' ' '{print $2}'")
    gateway2 = gateway2_origin.read().replace('\n','')
    #version变量生成
    version_stdin, version_origin, version_stderr = ssh_db2.exec_command("source /home/oracle/.profile;sqlplus -v |awk -F ' ' '{print $3}'")
    version = version_origin.read()
    #os_version变量生成
    os_version_stdin, os_version_origin, os_version_stderr = ssh_db2.exec_command("uname;oslevel -s")
    #os_version = os_version_origin.read().split('\n')
    os_version = os_version_origin.read()
    print os_version
    #timezone2变量生成
    timezone2_stdin, timezone2_origin, timezone2_stderr = ssh_db2.exec_command("date '+%Y-%m-%d %H:%M:%S';echo $TZ")
    timezone2 = timezone2_origin.read()
    #syslog2变量生成
    syslog2_stdin, syslog2_origin, syslog2_stderr = ssh_db2.exec_command("errpt")
    syslog2 = syslog2_origin.read()
    #filesystem2变量生成
    filesystem2_stdin, filesystem2_origin, filesystem2_stderr = ssh_db2.exec_command("df -g")
    filesystem2 = filesystem2_origin.read().split()
    #cluster_alert2变量生成
    cmd_cluster_alert2_1 = "tail -200 %s/log/" % grid_home
    cmd_cluster_alert2_2 = "%s/" % hostname2
    cmd_cluster_alert2_3 = "alert%s.log" % hostname2
    cmd_cluster_alert2 = cmd_cluster_alert2_1 + cmd_cluster_alert2_2 + cmd_cluster_alert2_3
    print cmd_cluster_alert2
    cluster_alert2_stdin, cluster_alert2_origin, cluster_alert2_stderr = ssh_db2.exec_command(cmd_cluster_alert2)
    cluster_alert2 = cluster_alert2_origin.read()
    #crsdlog2变量生成
    cmd_crsdlog2_1 = "tail -200 %s" % grid_home
    cmd_crsdlog2_2 = "/log/%s/crsd/crsd.log" % hostname2
    cmd_crsdlog2 = cmd_crsdlog2_1 + cmd_crsdlog2_2
    print cmd_crsdlog2
    crsdlog2_stdin, crsdlog2_origin, crsdlog2_stderr = ssh_db2.exec_command(cmd_crsdlog2)
    crsdlog2 = crsdlog2_origin.read()
    #cssdlog2变量生成
    cmd_cssdlog2_1 = "tail -200 %s" % grid_home
    cmd_cssdlog2_2 = "/log/%s/cssd/ocssd.log" % hostname2
    cmd_cssdlog2 = cmd_cssdlog2_1 + cmd_cssdlog2_2
    print cmd_cssdlog2
    cssdlog2_stdin, cssdlog2_origin, cssdlog2_stderr = ssh_db2.exec_command(cmd_cssdlog2)
    cssdlog2 = cssdlog2_origin.read()
    #asm_alert2变量生成
    cmd_asm_alert2 = "tail -200 %s/diag/asm/+asm/+ASM2/trace/alert_+ASM2.log|grep ORA-" % grid_base
    asm_alert2_stdin, asm_alert2_origin, asm_alert2_stderr = ssh_db2.exec_command(cmd_asm_alert2)
    asm_alert2 = asm_alert2_origin.read()
    #listener2变量生成
    listener2_stdin, listener2_origin, listener2_stderr = ssh_db2.exec_command("source /home/grid/.profile;lsnrctl stat")
    listener2 = listener2_origin.read()
    #listener_log2变量生成
    cmd_listener_log2_1 = "tail -1000 %s/diag/tnslsnr/" % grid_base
    cmd_listener_log2_2 = "%s/listener/trace/listener.log|grep TNS-" % hostname2
    cmd_listener_log2 = cmd_listener_log2_1 + cmd_listener_log2_2
    print cmd_listener_log2
    listener_log2_stdin, listener_log2_origin, listener_log2_stderr = ssh_db2.exec_command(cmd_listener_log2)
    listener_log2 = listener_log2_origin.read()
    #ssh_db2连接关闭
    ssh_db2.close()
#word路径，函数定义
    doc_path = r'check_database.docx'
    new_doc(userinfo,host,school_name,author,frequency,ip1,ip2,port1,port2,username1,username2,password1,password2,time,db_name,db_name_output,oracle_base,cpu1,cpu2,memory1,memory2,swap1,swap2,gateway1,gateway2,version,os_version,timezone1,timezone2,syslog1,syslog2,filesystem1,filesystem2,cluster,cluster_alert1,cluster_alert2,crsdlog1,crsdlog2,cssdlog1,cssdlog2,ocr_backup,asm_sql,asm_alert1,asm_alert2,database_version_sql,sga_sql,controlfile_sql,redolog_sql,tablespace_sql,virus_sql,listener1,listener2,listener_log1,listener_log2,doc_path)